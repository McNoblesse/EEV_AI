"""
Document Processing Utilities for Knowledge Base
Handles extraction, chunking, and embedding of various document types
"""

import io
import os
import logging
import csv
import tempfile
from typing import Tuple, List, Dict, Any, Optional
from fastapi import UploadFile
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

# Configure Tesseract path (Windows-specific, adjust if needed)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class DocumentProcessor:
    """Process various document types for knowledge base ingestion"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.csv', '.docx', '.doc'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        enable_ocr: bool = True
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_ocr = enable_ocr
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def process_document(
        self,
        file: UploadFile,
        category: str,
        enable_ocr: Optional[bool] = None
    ) -> Dict[str, Any]:
        """
        Process uploaded document and extract text
        
        Args:
            file: Uploaded file object
            category: Document category
            enable_ocr: Override OCR setting for this document
            
        Returns:
            Dictionary containing extracted text, chunks, and metadata
        """
        try:
            # Validate file
            self._validate_file(file)
            
            # Read file content
            content = await file.read()
            file_extension = os.path.splitext(file.filename)[1].lower()
            
            # Determine if OCR should be used
            use_ocr = enable_ocr if enable_ocr is not None else self.enable_ocr
            
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text, metadata = await self._extract_pdf(content, use_ocr)
            elif file_extension == '.txt':
                extracted_text, metadata = await self._extract_txt(content)
            elif file_extension == '.csv':
                extracted_text, metadata = await self._extract_csv(content)
            elif file_extension in ['.docx', '.doc']:
                extracted_text, metadata = await self._extract_docx(content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Chunk the text
            chunks = self._chunk_text(extracted_text)
            
            # Calculate statistics
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)
            
            result = {
                "text": extracted_text,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "word_count": word_count,
                "char_count": char_count,
                "file_type": file_extension,
                "category": category,
                "metadata": metadata,
                "preview": extracted_text[:500] if len(extracted_text) > 500 else extracted_text
            }
            
            logger.info(
                f"Processed document: {file.filename}, "
                f"chunks: {len(chunks)}, words: {word_count}"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file.filename}: {str(e)}")
            raise
    
    def _validate_file(self, file: UploadFile) -> None:
        """Validate file size and type"""
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_extension}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Note: file.size might not be available in all cases
        # Size check can be done after reading content if needed
    
    async def _extract_pdf(
        self,
        content: bytes,
        use_ocr: bool = True
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file
        Falls back to OCR for scanned PDFs
        """
        metadata = {
            "extraction_method": "text",
            "page_count": 0,
            "ocr_used": False
        }
        
        try:
            # Try standard text extraction first
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            metadata["page_count"] = len(pdf_reader.pages)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # If no text extracted and OCR is enabled, try OCR
            if not text.strip() and use_ocr:
                logger.info("No text found in PDF, attempting OCR...")
                text, ocr_metadata = await self._ocr_pdf(content)
                metadata.update(ocr_metadata)
                metadata["extraction_method"] = "ocr"
                metadata["ocr_used"] = True
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _ocr_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Perform OCR on scanned PDF
        """
        metadata = {
            "ocr_confidence": 0.0,
            "ocr_pages_processed": 0
        }
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(content, dpi=300)
            metadata["ocr_pages_processed"] = len(images)
            
            text = ""
            total_confidence = 0.0
            
            for i, image in enumerate(images):
                logger.info(f"OCR processing page {i+1}/{len(images)}...")
                
                # Perform OCR with confidence data
                try:
                    ocr_data = pytesseract.image_to_data(
                        image,
                        output_type=pytesseract.Output.DICT
                    )
                    
                    # Extract text and calculate average confidence
                    page_text = pytesseract.image_to_string(image)
                    text += page_text + "\n"
                    
                    # Calculate confidence
                    confidences = [
                        int(conf) for conf in ocr_data['conf']
                        if conf != '-1'
                    ]
                    if confidences:
                        page_confidence = sum(confidences) / len(confidences)
                        total_confidence += page_confidence
                        
                except Exception as ocr_error:
                    logger.warning(f"OCR failed for page {i+1}: {str(ocr_error)}")
                    continue
            
            # Calculate average confidence across all pages
            if metadata["ocr_pages_processed"] > 0:
                metadata["ocr_confidence"] = total_confidence / metadata["ocr_pages_processed"]
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"OCR processing failed: {str(e)}")
            raise ValueError(f"OCR failed: {str(e)}")
    
    async def _extract_txt(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with supported encodings")
            
            metadata = {
                "encoding": encoding_used,
                "line_count": len(text.split('\n'))
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT file: {str(e)}")
    
    async def _extract_csv(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file"""
        try:
            # Decode content
            text_content = content.decode('utf-8')
            
            # Parse CSV
            csv_reader = csv.DictReader(io.StringIO(text_content))
            rows = list(csv_reader)
            
            # Convert to text format
            text_parts = []
            for i, row in enumerate(rows):
                row_text = f"Record {i+1}:\n"
                for key, value in row.items():
                    row_text += f"  {key}: {value}\n"
                text_parts.append(row_text)
            
            text = "\n".join(text_parts)
            
            metadata = {
                "row_count": len(rows),
                "column_count": len(rows[0].keys()) if rows else 0,
                "columns": list(rows[0].keys()) if rows else []
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from CSV file: {str(e)}")
    
    async def _extract_docx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            # Save to temporary file (python-docx requires file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                # Open and extract text
                doc = Document(tmp_path)
                
                paragraphs = [para.text for para in doc.paragraphs]
                text = "\n".join(paragraphs)
                
                # Extract table content
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        table_text.append(row_text)
                
                if table_text:
                    text += "\n\nTables:\n" + "\n".join(table_text)
                
                metadata = {
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables)
                }
                
                return text, metadata
                
            finally:
                # Clean up temporary file
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
                    
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks for embedding"""
        try:
            chunks = self.text_splitter.split_text(text)
            return chunks
        except Exception as e:
            logger.error(f"Text chunking failed: {str(e)}")
            raise ValueError(f"Failed to chunk text: {str(e)}")
    
    def estimate_tokens(self, text: str) -> int:
        """
        Estimate token count (rough approximation)
        ~4 characters per token on average
        """
        return len(text) // 4


# Singleton instance
document_processor = DocumentProcessor(
    chunk_size=1000,
    chunk_overlap=200,
    enable_ocr=True
)
